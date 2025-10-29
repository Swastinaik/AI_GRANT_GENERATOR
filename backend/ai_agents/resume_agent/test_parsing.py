"""
resume_generator.py

Improved resume generator using python-docx with two templates: 'classic' and 'modern'.
This version supports input data where `contact` is a list and experience items
use 'duration' and 'description'. Education uses 'year'.
"""
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import random
import string

# ----------------------- Helpers -----------------------

def _set_margins(doc, left=Inches(0.7), right=Inches(0.7), top=Inches(0.7), bottom=Inches(0.7)):
    section = doc.sections[0]
    section.left_margin = left
    section.right_margin = right
    section.top_margin = top
    section.bottom_margin = bottom

def _add_horizontal_rule(paragraph, size=12, color='000000'):
    """Add a horizontal rule (bottom border) to the paragraph.
    size is in eights of a point (w:sz uses eighths); larger value = thicker line.
    """
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')

    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)

    pBdr.append(bottom)
    pPr.append(pBdr)

def _set_paragraph_font(paragraph, name='Calibri', size=11, bold=False, italic=False):
    for run in paragraph.runs:
        r = run.font
        r.name = name
        r.size = Pt(size)
        r.bold = bold
        r.italic = italic

def _shade_paragraph(paragraph, fill_hex='F2F2F2'):
    """Apply a background shade to the paragraph. fill_hex e.g. 'F2F2F2'"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)

# ----------------------- Utilities for new data shape -----------------------

def parse_contact_list(contact_list):
    """Take a list like ['email', 'phone', 'linkedin', 'city'] and map to dict keys.
    Uses simple heuristics; unknown items become part of 'other' or 'location'.
    """
    contact = {'email': None, 'phone': None, 'linkedin': None, 'location': None, 'other': []}
    for item in contact_list:
        if not item or not isinstance(item, str):
            continue
        s = item.strip()
        if '@' in s and contact['email'] is None:
            contact['email'] = s
        elif re.search(r'\+?\d[\d\s-]{5,}', s) and contact['phone'] is None:
            contact['phone'] = s
        elif 'linkedin' in s.lower() or s.lower().startswith('http') and 'linkedin' in s.lower():
            contact['linkedin'] = s
        else:
            # if it looks like 'City, Country' or contains letters only, treat as location
            if re.search(r'[A-Za-z]', s) and (',' in s or len(s.split()) <= 3):
                if contact['location'] is None:
                    contact['location'] = s
                else:
                    contact['other'].append(s)
            else:
                contact['other'].append(s)
    # collapse other into a single location string if no explicit location
    if not contact['location'] and contact['other']:
        contact['location'] = ', '.join(contact['other'])
    return contact

def split_description_to_bullets(description):
    """Split a free-form description into 1-4 bullets. Uses punctuation as split points."""
    if not description:
        return []
    # Normalize whitespace
    text = re.sub(r'\s+', ' ', description.strip())
    # Split by sentence-ending punctuation, but keep meaningful clauses together
    parts = re.split(r'(?<=[.!?])\s+', text)
    bullets = [p.strip().rstrip('.').strip() for p in parts if p.strip()]
    # If too many tiny bullets, try splitting by semicolon or ' and ' as fallback
    if len(bullets) > 6:
        # join then split by semicolon
        bullets = [b.strip().rstrip('.').strip() for b in re.split(r';\s*', text) if b.strip()]
    # limit number of bullets for neatness
    if len(bullets) > 6:
        bullets = bullets[:6]
    return bullets

# ----------------------- Header / Contact -----------------------

def _add_contact_block_fullwidth(document, contact_dict, template='classic'):
    # Name (centered)
    name_p = document.add_paragraph()
    name_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    name_run = name_p.add_run(contact_dict.get('name', ''))
    if template == 'classic':
        name_run.font.name = 'Times New Roman'
        name_run.font.size = Pt(20)
        name_run.bold = True
    else:
        name_run.font.name = 'Calibri'
        name_run.font.size = Pt(22)
        name_run.bold = True

    # Title (centered)
    if contact_dict.get('title'):
        title_p = document.add_paragraph()
        title_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        t_run = title_p.add_run(contact_dict.get('title'))
        t_run.font.name = 'Calibri' if template == 'modern' else 'Times New Roman'
        t_run.font.size = Pt(11)
        t_run.italic = True

    # Contact line: join available parts
    contact_items = []
    for k in ('email', 'phone', 'linkedin', 'location'):
        if contact_dict.get(k):
            contact_items.append(contact_dict[k])
    if contact_items:
        c_p = document.add_paragraph(' \u00A0 | \u00A0 '.join(contact_items))
        c_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        _set_paragraph_font(c_p, name=('Calibri' if template == 'modern' else 'Times New Roman'), size=9)

    # rule
    hr = document.add_paragraph()
    _add_horizontal_rule(hr, size=14, color='808080')

# ----------------------- Section helpers -----------------------

def _make_heading(doc_or_cell, text, template='classic', shaded=False):
    # doc_or_cell may be Document or a cell
    p = doc_or_cell.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri' if template == 'modern' else 'Times New Roman'
    if shaded:
        _shade_paragraph(p, fill_hex='F3F6FB')
    return p

def _add_section(document_or_cell, title, body_renderer, template='classic', shaded=False):
    p = _make_heading(document_or_cell, title, template=template, shaded=shaded)
    body_renderer(document_or_cell)
    # add rule after section
    hr = document_or_cell.add_paragraph()
    _add_horizontal_rule(hr, size=8, color='D9D9D9')

# ----------------------- Renderers for content -----------------------

def _render_summary(container, summary_text, template='classic'):
    p = container.add_paragraph(summary_text)
    _set_paragraph_font(p, name=('Calibri' if template == 'modern' else 'Times New Roman'), size=10)

def _render_skills(container, skills_list, template='classic'):
    # skills as small tag-like lines
    if not skills_list:
        return
    # render inline top line
    top = container.add_paragraph(', '.join(skills_list))
    _set_paragraph_font(top, name=('Calibri' if template == 'modern' else 'Times New Roman'), size=10)
    # and a simple bulleted set below (limit for readability)
    for s in skills_list:
        p = container.add_paragraph(s, style='List Bullet')
        _set_paragraph_font(p, name=('Calibri' if template == 'modern' else 'Times New Roman'), size=10)

def _render_experience(container, experiences, template='classic'):
    for exp in experiences:
        title = exp.get('title') or ''
        company = exp.get('company') or ''
        location = exp.get('location') or ''
        # header
        p = container.add_paragraph()
        run = p.add_run(f"{title}")
        run.bold = True
        run.font.size = Pt(11)
        _set_paragraph_font(p, name=('Calibri' if template == 'modern' else 'Times New Roman'), size=11)

        # meta line: company and location
        meta = container.add_paragraph()
        meta_text = f"{company}" + (f" — {location}" if location else '')
        meta_run = meta.add_run(meta_text)
        meta_run.italic = True
        _set_paragraph_font(meta, name=('Calibri' if template == 'modern' else 'Times New Roman'), size=10)

        # duration (use 'duration' or fallback to 'dates')
        duration = exp.get('duration') or exp.get('dates')
        if duration:
            dates_p = container.add_paragraph(duration)
            dates_p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            _set_paragraph_font(dates_p, name=('Calibri' if template == 'modern' else 'Times New Roman'), size=9)

        # description -> bullets
        bullets = []
        # first prefer explicit 'responsibilities' list (backwards compatibility)
        if exp.get('responsibilities') and isinstance(exp.get('responsibilities'), list):
            bullets = [b for b in exp.get('responsibilities') if b]
        else:
            # split description paragraphs into bullets
            desc = exp.get('description') or exp.get('summary') or ''
            bullets = split_description_to_bullets(desc)

        # render bullets
        for bullet in bullets:
            b = container.add_paragraph(bullet, style='List Bullet')
            _set_paragraph_font(b, name=('Calibri' if template == 'modern' else 'Times New Roman'), size=10)

def _render_education(container, education_list, template='classic'):
    for ed in education_list:
        p = container.add_paragraph()
        p.add_run(f"{ed.get('degree', '')}").bold = True
        if ed.get('institution'):
            p.add_run(f", {ed.get('institution')}")
        _set_paragraph_font(p, name=('Calibri' if template == 'modern' else 'Times New Roman'), size=11)
        year = ed.get('year') or ed.get('dates') or ed.get('duration')
        if year or ed.get('location'):
            info = container.add_paragraph(f"{year or ''}" + (f" — {ed.get('location')}" if ed.get('location') else ''))
            _set_paragraph_font(info, name=('Calibri' if template == 'modern' else 'Times New Roman'), size=10)

def _render_projects(container, projects, template='classic'):
    for proj in projects:
        p = container.add_paragraph()
        p.add_run(proj.get('name', '')).bold = True
        if proj.get('description'):
            d = container.add_paragraph(proj.get('description'))
            _set_paragraph_font(d, name=('Calibri' if template == 'modern' else 'Times New Roman'), size=10)

def _render_certifications(container, certs, template='classic'):
    for c in certs:
        p = container.add_paragraph(c)
        _set_paragraph_font(p, name=('Calibri' if template == 'modern' else 'Times New Roman'), size=10)

# ----------------------- Main generator -----------------------

def generate_resume_docx(data, template='classic'):
    """Generates a .docx resume using the provided data dict.

    template: 'classic' or 'modern'
    data fields: name, title, contact (list or dict), summary, experience (list), education (list), skills (list), projects (list), certifications (list)
    """
    # Normalize contact input: accept list or dict
    directory = "resume"
    os.makedirs(directory, exist_ok=True)
    characters = string.ascii_letters  + string.digits 
    file_name = 'resume-' +''.join(random.choices(characters, k=5)) + 'resume.docx'
    output_path = os.path.join(directory, file_name)
    raw_contact = data.get('contact', {})
    if isinstance(raw_contact, list):
        parsed = parse_contact_list(raw_contact)
        # include top-level name/title if provided in data
        parsed['name'] = data.get('name') or parsed.get('name')
        parsed['title'] = data.get('title') or parsed.get('title')
        contact_dict = parsed
    elif isinstance(raw_contact, dict):
        contact_dict = raw_contact.copy()
        # ensure name and title from top-level keys
        if data.get('name'):
            contact_dict['name'] = data.get('name')
        if data.get('title'):
            contact_dict['title'] = data.get('title')
    else:
        # fallback to top-level name/title only
        contact_dict = {'name': data.get('name'), 'title': data.get('title'), 'email': None, 'phone': None, 'linkedin': None, 'location': None}

    doc = Document()
    _set_margins(doc)

    # default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman' if template == 'classic' else 'Calibri'
    style.font.size = Pt(11)

    # --- Modern template uses a 2-column table layout ---
    if template == 'modern':
        # Create a wide table with two columns: left sidebar and main content
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        # attempt to set column widths
        try:
            table.columns[0].width = Inches(2.1)
            table.columns[1].width = Inches(4.5)
        except Exception:
            pass
        left, right = table.rows[0].cells

        # Left: contact + skills + certs in a shaded box
        c_para = left.add_paragraph()
        c_para.add_run(contact_dict.get('name', '') or '').bold = True
        _set_paragraph_font(c_para, name='Calibri', size=14)
        if contact_dict.get('title'):
            t = left.add_paragraph(contact_dict.get('title'))
            _set_paragraph_font(t, name='Calibri', size=10, italic=True)

        # small contact details
        contact_items = []
        for k in ('email', 'phone', 'linkedin', 'location'):
            if contact_dict.get(k):
                contact_items.append(contact_dict[k])
        if contact_items:
            left.add_paragraph('\n'.join(contact_items))

        # separator
        left.add_paragraph()
        _add_horizontal_rule(left.add_paragraph(), size=8, color='D9D9D9')

        # Skills
        if data.get('skills'):
            _add_section(left, 'Skills', lambda c: _render_skills(c, data['skills'], template=template), template=template, shaded=False)

        # Certifications
        if data.get('certifications'):
            _add_section(left, 'Certifications', lambda c: _render_certifications(c, data['certifications'], template=template), template=template, shaded=False)

        # Right: main content
        if data.get('summary'):
            _add_section(right, 'Summary', lambda c: _render_summary(c, data['summary'], template=template), template=template, shaded=True)

        if data.get('experience'):
            _add_section(right, 'Experience', lambda c: _render_experience(c, data['experience'], template=template), template=template)

        if data.get('projects'):
            _add_section(right, 'Projects', lambda c: _render_projects(c, data['projects'], template=template), template=template)

        if data.get('education'):
            _add_section(right, 'Education', lambda c: _render_education(c, data['education'], template=template), template=template)

    else:
        # Classic template: full-width header + sections
        _add_contact_block_fullwidth(doc, contact_dict, template=template)

        if data.get('summary'):
            _add_section(doc, 'Summary', lambda d: _render_summary(d, data['summary'], template=template), template=template)

        if data.get('experience'):
            _add_section(doc, 'Experience', lambda d: _render_experience(d, data['experience'], template=template), template=template)

        if data.get('education'):
            _add_section(doc, 'Education', lambda d: _render_education(d, data['education'], template=template), template=template)

        if data.get('projects'):
            _add_section(doc, 'Projects', lambda d: _render_projects(d, data['projects'], template=template), template=template)

        if data.get('skills'):
            _add_section(doc, 'Skills', lambda d: _render_skills(d, data['skills'], template=template), template=template)

        if data.get('certifications'):
            _add_section(doc, 'Certifications', lambda d: _render_certifications(d, data['certifications'], template=template), template=template)

    # Save
    doc.save(output_path)
    absolute_path = os.path.abspath(output_path)
    return absolute_path
