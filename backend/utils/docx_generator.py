import io
from docx import Document
from fastapi.responses import StreamingResponse

def generate_docx_stream(data: dict) -> io.BytesIO:
    doc = Document()
    doc.add_heading('Grant Proposal', 0)
    for title, content in data.items():
        doc.add_heading(title, level=1)
        doc.add_paragraph(str(content) if content else '')
    
    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream

